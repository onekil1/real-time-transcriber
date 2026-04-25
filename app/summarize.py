from __future__ import annotations

import json
import re
import time
from typing import Any

import requests

from .config import (
    LLM_CHUNK_CHARS,
    LLM_MAX_TOKENS,
    LLM_TEMPERATURE,
    LM_STUDIO_API_KEY,
    LM_STUDIO_BASE_URL,
    LM_STUDIO_MODEL,
    PROMPTS_DIR,
)


class LLMUnavailable(RuntimeError):
    pass


# ---------- шаблоны отчётов ---------------------------------------------------
# Секция: (json_key, heading, kind). kind = "text" | "list".
TEMPLATES: dict[str, dict[str, Any]] = {
    "meeting": {
        "label": "Совещание",
        "prompt_file": "meeting_ru.txt",
        "sections": [
            ("agenda", "Повестка", "list"),
            ("summary", "Резюме", "text"),
            ("decisions", "Решения", "list"),
            ("open_questions", "Открытые вопросы", "list"),
        ],
    },
    "presentation": {
        "label": "Выступление",
        "prompt_file": "presentation_ru.txt",
        "sections": [
            ("participants", "Участники", "list"),
            ("key_theses", "Ключевые тезисы", "list"),
            ("summary", "Резюме", "text"),
            ("questions_raised", "Вопросы и обсуждение", "list"),
            ("conclusions", "Выводы", "list"),
        ],
    },
    "brainstorm": {
        "label": "Мозговой штурм",
        "prompt_file": "brainstorm_ru.txt",
        "sections": [
            ("topic", "Тема", "text"),
            ("summary", "Резюме", "text"),
            ("ideas", "Все идеи", "list"),
            ("grouped_themes", "Сгруппированные направления", "list"),
            ("selected_ideas", "Наиболее перспективные", "list"),
            ("next_steps", "Следующие шаги", "list"),
        ],
    },
    "planning": {
        "label": "Планирование задач",
        "prompt_file": "planning_ru.txt",
        "sections": [
            ("goals", "Цели", "list"),
            ("summary", "Резюме", "text"),
            ("tasks", "Задачи", "list"),
            ("priorities", "Приоритеты", "list"),
            ("blockers", "Блокеры и риски", "list"),
        ],
    },
}

DEFAULT_TEMPLATE = "meeting"


def list_templates() -> list[dict[str, str]]:
    return [{"key": k, "label": v["label"]} for k, v in TEMPLATES.items()]


def _template(key: str | None) -> dict[str, Any]:
    return TEMPLATES.get(key or "") or TEMPLATES[DEFAULT_TEMPLATE]


def _empty_report(tmpl: dict[str, Any]) -> dict[str, Any]:
    out: dict[str, Any] = {"title": "", "date": ""}
    for key, _h, kind in tmpl["sections"]:
        out[key] = "" if kind == "text" else []
    return out


# ---------- LM Studio ---------------------------------------------------------

def _resolve_model() -> str:
    if LM_STUDIO_MODEL:
        return LM_STUDIO_MODEL
    try:
        r = requests.get(
            f"{LM_STUDIO_BASE_URL}/models",
            headers={"Authorization": f"Bearer {LM_STUDIO_API_KEY}"},
            timeout=5,
        )
        r.raise_for_status()
        data = r.json().get("data") or []
        if data and data[0].get("id"):
            return data[0]["id"]
    except requests.RequestException as e:
        raise LLMUnavailable(
            f"LM Studio недоступен на {LM_STUDIO_BASE_URL}: {e}"
        ) from e
    raise LLMUnavailable("В LM Studio не загружено ни одной модели")


def ping() -> dict[str, Any]:
    base: dict[str, Any] = {
        "base_url": LM_STUDIO_BASE_URL,
        "configured_model": LM_STUDIO_MODEL or None,
    }
    try:
        r = requests.get(
            f"{LM_STUDIO_BASE_URL}/models",
            headers={"Authorization": f"Bearer {LM_STUDIO_API_KEY}"},
            timeout=3,
        )
        r.raise_for_status()
        models = [m.get("id") for m in r.json().get("data", []) if m.get("id")]
        if not models:
            return {
                **base,
                "ok": False,
                "models": [],
                "error": "В LM Studio не загружено ни одной модели",
            }
        active = LM_STUDIO_MODEL if LM_STUDIO_MODEL in models else models[0]
        return {**base, "ok": True, "models": models, "active_model": active}
    except requests.RequestException as e:
        return {**base, "ok": False, "error": str(e)}


def _chat(prompt: str, max_tokens: int = LLM_MAX_TOKENS) -> str:
    model = _resolve_model()
    body = {
        "model": model,
        "messages": [
            {"role": "system", "content": "/no_think"},
            {"role": "user", "content": prompt},
        ],
        "temperature": LLM_TEMPERATURE,
        "max_tokens": max_tokens,
        "stream": False,
    }
    headers = {
        "Authorization": f"Bearer {LM_STUDIO_API_KEY}",
        "Content-Type": "application/json",
    }
    url = f"{LM_STUDIO_BASE_URL}/chat/completions"

    # 1 ретрай на сетевые сбои / залипания при первом промпте после загрузки
    # модели в LM Studio. На HTTP-ошибки (4xx/5xx) не ретраим — обычно это
    # некорректный запрос или OOM, повтор не поможет.
    last_err: Exception | None = None
    for attempt in (0, 1):
        try:
            r = requests.post(url, headers=headers, json=body, timeout=600)
            r.raise_for_status()
            msg = r.json()["choices"][0]["message"]
            return (msg.get("content") or "").strip()
        except (requests.ConnectionError, requests.Timeout) as e:
            last_err = e
            if attempt == 0:
                time.sleep(2.0)
                continue
            raise LLMUnavailable(f"LM Studio недоступна: {e}") from e
        except requests.RequestException as e:
            raise LLMUnavailable(f"LM Studio HTTP error: {e}") from e
    raise LLMUnavailable(f"LM Studio недоступна: {last_err}")


# ---------- JSON-парсер -------------------------------------------------------

_JSON_RE = re.compile(r"\{.*\}", re.DOTALL)


def _parse_json(raw: str) -> dict[str, Any]:
    raw = raw.strip()
    if raw.startswith("```"):
        raw = re.sub(r"^```(?:json)?\s*|\s*```$", "", raw, flags=re.IGNORECASE)
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        m = _JSON_RE.search(raw)
        if m:
            try:
                return json.loads(m.group(0))
            except json.JSONDecodeError:
                pass
    return {}


_SENT_SPLIT_RE = re.compile(r"(?<=[.!?…])\s+")


def _limit_sentences(text: str, max_sentences: int = 5) -> str:
    text = (text or "").strip()
    if not text:
        return ""
    parts = _SENT_SPLIT_RE.split(text)
    return " ".join(parts[:max_sentences]).strip()


# ---------- map-reduce --------------------------------------------------------

def _chunk_text(text: str, size: int) -> list[str]:
    if len(text) <= size:
        return [text]
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + size, len(text))
        if end < len(text):
            nl = text.rfind("\n", start + size // 2, end)
            dot = text.rfind(". ", start + size // 2, end)
            cut = max(nl, dot)
            if cut > start:
                end = cut + 1
        chunks.append(text[start:end])
        start = end
    return chunks


def _merge_reports(parts: list[dict[str, Any]], tmpl: dict[str, Any]) -> dict[str, Any]:
    merged = _empty_report(tmpl)
    text_buffers: dict[str, list[str]] = {}
    for p in parts:
        if p.get("title") and not merged["title"]:
            merged["title"] = p["title"]
        if p.get("date") and not merged["date"]:
            merged["date"] = p["date"]
        for key, _h, kind in tmpl["sections"]:
            val = p.get(key)
            if kind == "list":
                for item in val or []:
                    if item and item not in merged[key]:
                        merged[key].append(item)
            else:
                if val:
                    text_buffers.setdefault(key, []).append(str(val))
    for key, _h, kind in tmpl["sections"]:
        if kind == "text" and key in text_buffers:
            joined = " ".join(text_buffers[key])
            if key == "summary":
                merged[key] = _limit_sentences(joined)
            elif key == "topic":
                merged[key] = _limit_sentences(joined, 1)
            else:
                merged[key] = joined
    return merged


def _to_markdown(r: dict[str, Any], tmpl: dict[str, Any]) -> str:
    def lst(items: list[Any]) -> str:
        return "\n".join(f"- {x}" for x in items) if items else "_—_"

    lines = [
        f"# {r.get('title') or 'Отчёт'}",
        f"**Тип:** {tmpl['label']}  ",
        f"**Дата:** {r.get('date') or '—'}",
        "",
    ]
    for key, heading, kind in tmpl["sections"]:
        lines.append(f"## {heading}")
        if kind == "list":
            lines.append(lst(r.get(key) or []))
        else:
            val = r.get(key) or ""
            if key == "summary":
                val = _limit_sentences(val)
            elif key == "topic":
                val = _limit_sentences(val, 1)
            lines.append(val or "_—_")
        lines.append("")
    return "\n".join(lines).rstrip() + "\n"


def generate_report(
    transcript: str, progress_cb=None, template: str | None = None
) -> dict[str, Any]:
    tmpl = _template(template)
    prompt_template = (PROMPTS_DIR / tmpl["prompt_file"]).read_text(encoding="utf-8")
    chunks = _chunk_text(transcript, LLM_CHUNK_CHARS)
    if progress_cb:
        progress_cb("llm:start", 0.0)

    part_reports: list[dict[str, Any]] = []
    for i, chunk in enumerate(chunks):
        prompt = prompt_template.replace("{TRANSCRIPT}", chunk)
        raw = _chat(prompt)
        parsed = _parse_json(raw)
        base = _empty_report(tmpl)
        if parsed:
            base.update({k: parsed.get(k, base[k]) for k in base})
        part_reports.append(base)
        if progress_cb:
            progress_cb("llm:chunk", (i + 1) / len(chunks))

    merged = part_reports[0] if len(part_reports) == 1 else _merge_reports(part_reports, tmpl)
    merged["_template"] = tmpl["label"]
    md = _to_markdown(merged, tmpl)

    if progress_cb:
        progress_cb("llm:done", 1.0)

    return {"data": merged, "md": md, "template": template or DEFAULT_TEMPLATE}


# ---------- DOCX export ------------------------------------------------------

def render_docx(report: dict[str, Any] | None, template_key: str | None, title: str):
    """Собрать DOCX-отчёт. Возвращает Document (python-docx)."""
    from docx import Document
    from docx.shared import Pt

    tmpl = _template(template_key)
    data = (report or {}).get("data") or {}
    md_only = (report or {}).get("md") or ""

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading(data.get("title") or title or "Отчёт", level=1)
    info = doc.add_paragraph()
    info.add_run("Тип: ").bold = True
    info.add_run(tmpl["label"])
    info.add_run("    Дата: ").bold = True
    info.add_run(data.get("date") or "—")

    if not data and md_only:
        # нет структурированных данных — вставим markdown как есть
        for line in md_only.splitlines():
            doc.add_paragraph(line)
        return doc

    for key, heading, kind in tmpl["sections"]:
        doc.add_heading(heading, level=2)
        val = data.get(key)
        if kind == "list":
            items = val or []
            if not items:
                doc.add_paragraph("—")
            else:
                for it in items:
                    doc.add_paragraph(f"— {it}")
        else:
            text = str(val or "").strip()
            if key == "summary":
                text = _limit_sentences(text)
            elif key == "topic":
                text = _limit_sentences(text, 1)
            doc.add_paragraph(text or "—")

    return doc
